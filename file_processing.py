import os
import shutil
import json
import datetime
import re
from pathlib import Path

from config import MACHINE_READABLE_FORMATS, OUTPUT_FORMATS
from db_ops import insert_document


def sanitize_filename(filename: str) -> str:
    sanitized = re.sub(r'[<>:"/\\|?*]', '_', filename)
    sanitized = sanitized.strip('. ')
    return sanitized[:255]


def generate_unique_filename(base_name: str, storage_dir: str, extension: str) -> str:
    counter = 1
    original_name = base_name
    while os.path.exists(os.path.join(storage_dir, f"{base_name}.{extension}")):
        base_name = f"{original_name}_{counter}"
        counter += 1
    return f"{base_name}.{extension}"


def detect_file_readability(file_path: str) -> bool:
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    if file_ext in ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif', 'webp']:
        return False
    if file_ext == 'pdf':
        try:
            import pdfplumber
            with pdfplumber.open(file_path) as pdf:
                if len(pdf.pages) > 0:
                    text = pdf.pages[0].extract_text()
                    return bool(text and text.strip())
        except:
            pass
        return False
    return file_ext in MACHINE_READABLE_FORMATS


def extract_text_from_file(file_path: str, force_ocr: bool = False):
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    if force_ocr or not detect_file_readability(file_path):
        if file_ext == 'pdf':
            return ocr_pdf_to_text(file_path)
        elif file_ext in ['jpg', 'jpeg', 'png', 'tiff', 'bmp', 'gif', 'webp']:
            return ocr_image_to_text(file_path)
    if file_ext == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_ext == 'docx':
        return extract_text_from_docx(file_path)
    elif file_ext == 'doc':
        return extract_text_from_doc(file_path)
    elif file_ext == 'rtf':
        return extract_text_from_rtf(file_path)
    elif file_ext == 'odt':
        return extract_text_from_odt(file_path)
    elif file_ext == 'txt':
        return extract_text_from_txt(file_path)
    elif file_ext in ['csv', 'json', 'xml', 'html', 'md']:
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_ext}")


def extract_text_from_pdf(pdf_path: str):
    try:
        import pdfplumber
        with pdfplumber.open(pdf_path) as pdf:
            full_text = []
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    full_text.append(text.strip())
            joined = "\n".join(full_text).strip()
            return (len(joined) > 0), joined, "direct_extraction"
    except ImportError:
        raise RuntimeError("pdfplumber required: pip install pdfplumber")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from PDF: {e}")


def extract_text_from_docx(docx_path: str):
    try:
        from docx import Document
        doc = Document(docx_path)
        full_text = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text.strip())
        joined = "\n".join(full_text).strip()
        return (len(joined) > 0), joined, "direct_extraction"
    except ImportError:
        raise RuntimeError("python-docx required: pip install python-docx")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOCX: {e}")


def extract_text_from_doc(doc_path: str):
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.visible = False
        doc = word.Documents.Open(os.path.abspath(doc_path))
        text = doc.Range().Text
        doc.Close()
        word.Quit()
        return (len(text.strip()) > 0), text.strip(), "direct_extraction"
    except ImportError:
        raise RuntimeError("pywin32 required: pip install pywin32")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from DOC: {e}")


def extract_text_from_rtf(rtf_path: str):
    try:
        from striprtf.striprtf import rtf_to_text
        with open(rtf_path, 'r', encoding='utf-8') as f:
            rtf_content = f.read()
        text = rtf_to_text(rtf_content)
        return (len(text.strip()) > 0), text.strip(), "direct_extraction"
    except ImportError:
        raise RuntimeError("striprtf required: pip install striprtf")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from RTF: {e}")


def extract_text_from_odt(odt_path: str):
    try:
        from odf import text, teletype
        from odf.opendocument import load
        doc = load(odt_path)
        all_text = []
        for paragraph in doc.getElementsByType(text.P):
            all_text.append(teletype.extractText(paragraph))
        joined = "\n".join(all_text).strip()
        return (len(joined) > 0), joined, "direct_extraction"
    except ImportError:
        raise RuntimeError("odfpy required: pip install odfpy")
    except Exception as e:
        raise RuntimeError(f"Failed to extract text from ODT: {e}")


def extract_text_from_txt(file_path: str):
    try:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()
        return (len(text.strip()) > 0), text, "direct_read"
    except Exception as e:
        raise RuntimeError(f"Failed to read text file: {e}")


def ocr_image_to_text(image_path: str):
    try:
        import pytesseract
        from PIL import Image
        img = Image.open(image_path)
        text = pytesseract.image_to_string(img)
        return (len(text.strip()) > 0), text.strip(), "ocr"
    except ImportError:
        raise RuntimeError("pytesseract and pillow required: pip install pytesseract pillow")
    except Exception as e:
        raise RuntimeError(f"Failed to perform OCR on image: {e}")


def ocr_pdf_to_text(pdf_path: str):
    try:
        from pdf2image import convert_from_path
        import pytesseract
        images = convert_from_path(pdf_path, dpi=300)
        texts = []
        for i, img in enumerate(images):
            try:
                text = pytesseract.image_to_string(img)
                texts.append(text)
            except Exception as e:
                texts.append(f"[ERROR extracting page {i}: {e}]")
        full_text = "\n".join(texts).strip()
        return (len(full_text) > 0), full_text, "ocr"
    except ImportError:
        raise RuntimeError("pdf2image and pytesseract required: pip install pdf2image pytesseract")
    except Exception as e:
        raise RuntimeError(f"Failed to perform OCR on PDF: {e}")


def convert_to_output_format(text: str, output_format: str, output_path: str) -> None:
    if output_format == 'txt':
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    elif output_format == 'docx':
        try:
            from docx import Document
            doc = Document()
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph)
            doc.save(output_path)
        except ImportError:
            raise RuntimeError("python-docx required: pip install python-docx")
    elif output_format == 'html':
        html_content = f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="UTF-8">
    <title>Converted Document</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 40px; line-height: 1.6; }}
        p {{ margin-bottom: 10px; }}
    </style>
</head>
<body>
"""
        for paragraph in text.split('\n'):
            if paragraph.strip():
                html_content += f"    <p>{paragraph.strip()}</p>\n"
        html_content += "</body>\n</html>"
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(html_content)
    elif output_format == 'md':
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text)
    elif output_format == 'json':
        json_data = {
            "content": text,
            "paragraphs": [p.strip() for p in text.split('\n') if p.strip()],
            "converted_at": datetime.datetime.utcnow().isoformat(),
            "word_count": len(text.split())
        }
        with open(output_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, indent=2, ensure_ascii=False)
    elif output_format == 'pdf':
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph
            from reportlab.lib.styles import getSampleStyleSheet
            doc = SimpleDocTemplate(output_path, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            for paragraph in text.split('\n'):
                if paragraph.strip():
                    story.append(Paragraph(paragraph, styles['Normal']))
            doc.build(story)
        except ImportError:
            raise RuntimeError("reportlab required: pip install reportlab")


def process_file(file_path: str, output_format: str = 'txt', custom_name: str = "",
                 tags: str = "", description: str = "", force_ocr: bool = False) -> dict:
    storage_dir = os.path.join(os.path.dirname(__file__), "storage")
    os.makedirs(storage_dir, exist_ok=True)

    base_name = os.path.basename(file_path)
    file_ext = Path(file_path).suffix.lower().lstrip('.')
    file_size = os.path.getsize(file_path)

    is_machine_readable = detect_file_readability(file_path) and not force_ocr

    if custom_name:
        display_name = sanitize_filename(custom_name)
    else:
        display_name = Path(base_name).stem
        custom_name = base_name

    stored_filename = generate_unique_filename(display_name, storage_dir, file_ext)
    stored_path = os.path.join(storage_dir, stored_filename)
    shutil.copy2(file_path, stored_path)

    try:
        readable, extracted_text, method = extract_text_from_file(stored_path, force_ocr)
        word_count = len(extracted_text.split()) if extracted_text else 0
    except Exception as e:
        print(f"Text extraction failed: {e}")
        readable, extracted_text, method, word_count = False, "", "failed", 0

    extracted_text_path = ""
    output_path = ""

    if readable and extracted_text:
        extracted_filename = f"{display_name}_extracted.txt"
        extracted_text_path = os.path.join(storage_dir, extracted_filename)
        with open(extracted_text_path, "w", encoding="utf-8") as f:
            f.write(extracted_text)

        output_filename = f"{display_name}_converted.{output_format}"
        output_path = os.path.join(storage_dir, output_filename)
        try:
            convert_to_output_format(extracted_text, output_format, output_path)
            print(f"Successfully converted to {output_format.upper()}: {output_path}")
        except Exception as e:
            print(f"Conversion to {output_format} failed: {e}")
            output_path = ""

    doc_id = insert_document(base_name, custom_name, stored_path, file_ext, is_machine_readable,
                             readable, extracted_text_path, output_format, output_path, method,
                             file_size, word_count, tags, description, extracted_text)

    return {
        "id": doc_id,
        "name": base_name,
        "custom_name": custom_name,
        "stored_path": stored_path,
        "original_format": file_ext,
        "is_machine_readable": is_machine_readable,
        "readable": readable,
        "extracted_text_path": extracted_text_path,
        "output_format": output_format,
        "output_path": output_path,
        "processing_method": method,
        "file_size": file_size,
        "word_count": word_count,
        "tags": tags,
        "description": description
    }


